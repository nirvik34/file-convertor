import json
from io import BytesIO
from flask import Flask
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import Workbook
from pptx import Presentation
from pydub import AudioSegment

app = Flask(__name__)
CORS(app)

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    output_type = request.form.get('output_type')
    filename = file.filename
    ext = filename.split('.')[-1].lower()

    file_stream = BytesIO(file.read())

    output_buffer = BytesIO()
    output_filename = f"converted.{output_type}"

    try:
        if ext == 'pdf' and output_type == 'txt':
            reader = PdfReader(file_stream)
            text = '\n'.join([page.extract_text() or '' for page in reader.pages])
            output_buffer.write(text.encode('utf-8'))

        elif ext == 'pdf' and output_type == 'docx':
            reader = PdfReader(file_stream)
            doc = Document()
            for page in reader.pages:
                doc.add_paragraph(page.extract_text() or '')
            doc.save(output_buffer)

        elif ext == 'docx' and output_type == 'txt':
            doc = Document(file_stream)
            for para in doc.paragraphs:
                output_buffer.write((para.text + '\n').encode('utf-8'))

        elif ext == 'txt' and output_type == 'docx':
            doc = Document()
            text = file_stream.read().decode('utf-8')
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(output_buffer)

        elif ext == 'json' and output_type == 'txt':
            data = json.load(file_stream)
            output_buffer.write(json.dumps(data, indent=4).encode('utf-8'))

        elif ext == 'mp3' and output_type == 'wav':
            sound = AudioSegment.from_file(file_stream, format='mp3')
            sound.export(output_buffer, format="wav")

        elif ext == 'wav' and output_type == 'mp3':
            sound = AudioSegment.from_file(file_stream, format='wav')
            sound.export(output_buffer, format="mp3")

        elif ext == 'mp4' and output_type == 'mp3':
            video = VideoFileClip(file_stream)
            audio_io = BytesIO()
            video.audio.write_audiofile(audio_io)
            audio_io.seek(0)
            return send_file(audio_io, as_attachment=True, download_name=output_filename, mimetype="audio/mpeg")

        elif ext == 'pptx' and output_type == 'pdf':
            return jsonify({'error': 'PPTX to PDF not supported in-memory'}), 400

        elif ext == 'xlsx' and output_type == 'csv':
            return jsonify({'error': 'XLSX to CSV not supported in-memory yet'}), 400

        else:
            return jsonify({'error': f'Conversion from .{ext} to .{output_type} not supported'}), 400

        output_buffer.seek(0)
        return send_file(output_buffer, as_attachment=True, download_name=output_filename)

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': 'Conversion failed'}), 500

if __name__ == "__main__":
    app.run(debug=True)
